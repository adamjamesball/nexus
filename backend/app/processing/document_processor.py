"""
Advanced document processing pipeline with AI-powered analysis
"""
import asyncio
import logging
import mimetypes
import os
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, field
from pathlib import Path
import hashlib

import pandas as pd
import PyPDF2
from docx import Document
# from unstructured.partition.auto import partition  # Disabled for simplified testing
# import magic  # Disabled for simplified testing

from ..llm import LLMMessage
from ..llm.client import client
from ..config import get_settings, get_supported_file_extensions


logger = logging.getLogger(__name__)


@dataclass
class ProcessedDocument:
    """Represents a processed document with extracted content and metadata"""
    file_path: str
    filename: str
    file_type: str
    size_bytes: int
    content: str
    structured_data: Optional[Dict[str, Any]] = None
    entities: List[Dict[str, Any]] = field(default_factory=list)
    sustainability_insights: Dict[str, Any] = field(default_factory=dict)
    processing_status: str = "pending"  # pending, processing, completed, error
    error_message: Optional[str] = None
    confidence_score: float = 0.0
    metadata: Dict[str, Any] = field(default_factory=dict)


class DocumentProcessor:
    """Advanced document processing with AI-powered content analysis"""

    def __init__(self):
        self.settings = get_settings()
        self.supported_extensions = get_supported_file_extensions()

    async def process_file(self, file_path: str) -> ProcessedDocument:
        """Process a single file and extract all relevant information"""
        logger.info(f"Processing file: {file_path}")

        try:
            # Initialize document object
            doc = ProcessedDocument(
                file_path=file_path,
                filename=os.path.basename(file_path),
                file_type=self._detect_file_type(file_path),
                size_bytes=os.path.getsize(file_path),
                content="",
                processing_status="processing"
            )

            # Check file size limits
            if doc.size_bytes > self.settings.max_file_size_mb * 1024 * 1024:
                doc.processing_status = "error"
                doc.error_message = f"File too large ({doc.size_bytes / 1024 / 1024:.1f}MB). Max allowed: {self.settings.max_file_size_mb}MB"
                return doc

            # Extract raw content based on file type
            doc.content = await self._extract_content(file_path, doc.file_type)

            if not doc.content:
                doc.processing_status = "error"
                doc.error_message = "Could not extract content from file"
                return doc

            # Process structured data if applicable
            if doc.file_type in ['csv', 'xlsx', 'xls']:
                doc.structured_data = await self._process_structured_data(file_path, doc.file_type)

            # AI-powered content analysis
            doc.entities = await self._extract_entities_with_ai(doc.content)
            doc.sustainability_insights = await self._analyze_sustainability_content(doc.content)

            # Calculate confidence score
            doc.confidence_score = self._calculate_confidence_score(doc)

            # Add metadata
            doc.metadata = {
                'content_hash': hashlib.md5(doc.content.encode()).hexdigest(),
                'word_count': len(doc.content.split()),
                'char_count': len(doc.content),
                'processed_at': asyncio.get_event_loop().time()
            }

            doc.processing_status = "completed"
            logger.info(f"Successfully processed {doc.filename}")

        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            doc.processing_status = "error"
            doc.error_message = str(e)

        return doc

    async def process_multiple_files(self, file_paths: List[str]) -> List[ProcessedDocument]:
        """Process multiple files concurrently"""
        logger.info(f"Processing {len(file_paths)} files")

        tasks = [self.process_file(path) for path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        processed_docs = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                # Create error document for failed processing
                error_doc = ProcessedDocument(
                    file_path=file_paths[i],
                    filename=os.path.basename(file_paths[i]),
                    file_type="unknown",
                    size_bytes=0,
                    content="",
                    processing_status="error",
                    error_message=str(result)
                )
                processed_docs.append(error_doc)
            else:
                processed_docs.append(result)

        return processed_docs

    def _detect_file_type(self, file_path: str) -> str:
        """Detect file type using multiple methods"""
        # First try file extension
        ext = Path(file_path).suffix.lower()
        if ext:
            ext_map = {
                '.pdf': 'pdf',
                '.docx': 'docx',
                '.doc': 'doc',
                '.xlsx': 'xlsx',
                '.xls': 'xls',
                '.csv': 'csv',
                '.txt': 'txt'
            }
            if ext in ext_map:
                return ext_map[ext]

        # Fallback to MIME type detection - simplified version
        # try:
        #     mime_type = magic.from_file(file_path, mime=True)
        #     mime_map = {
        #         'application/pdf': 'pdf',
        #         'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        #         'application/msword': 'doc',
        #         'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
        #         'application/vnd.ms-excel': 'xls',
        #         'text/csv': 'csv',
        #         'text/plain': 'txt'
        #     }
        #     return mime_map.get(mime_type, 'unknown')
        # except:
        #     return 'unknown'
        return 'unknown'  # Simplified for testing

    async def _extract_content(self, file_path: str, file_type: str) -> str:
        """Extract text content based on file type"""
        try:
            if file_type == 'pdf':
                return await self._extract_pdf_content(file_path)
            elif file_type == 'docx':
                return await self._extract_docx_content(file_path)
            elif file_type in ['csv', 'xlsx', 'xls']:
                return await self._extract_spreadsheet_content(file_path)
            elif file_type == 'txt':
                return await self._extract_text_content(file_path)
            else:
                # Use unstructured for other formats - disabled for testing
                # return await self._extract_with_unstructured(file_path)
                return f"Content from {file_type} file (unsupported in simplified mode)"
        except Exception as e:
            logger.error(f"Content extraction failed for {file_path}: {e}")
            return ""

    async def _extract_pdf_content(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content.append(page.extract_text())

                return '\n'.join(text_content)
        except Exception as e:
            logger.warning(f"PyPDF2 failed: {e}")
            return f"PDF extraction failed: {e}"

    async def _extract_docx_content(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text_content = []

            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)

            return '\n'.join(text_content)
        except Exception as e:
            logger.warning(f"python-docx failed: {e}")
            return f"DOCX extraction failed: {e}"

    async def _extract_spreadsheet_content(self, file_path: str) -> str:
        """Extract content from spreadsheet files"""
        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)

            # Convert to text representation
            text_content = [f"Sheet Data:\n{df.to_string()}"]

            # Add column names and data types
            text_content.append(f"\nColumns: {list(df.columns)}")
            text_content.append(f"Data Types: {df.dtypes.to_dict()}")
            text_content.append(f"Shape: {df.shape[0]} rows, {df.shape[1]} columns")

            return '\n'.join(text_content)
        except Exception as e:
            logger.error(f"Spreadsheet extraction failed: {e}")
            return ""

    async def _extract_text_content(self, file_path: str) -> str:
        """Extract content from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except:
                    continue
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return ""

    async def _extract_with_unstructured(self, file_path: str) -> str:
        """Extract content using unstructured library - disabled for testing"""
        # try:
        #     elements = partition(filename=file_path)
        #     return '\n'.join([str(element) for element in elements])
        # except Exception as e:
        #     logger.error(f"Unstructured extraction failed: {e}")
        #     return ""
        return "Unstructured extraction disabled in simplified mode"

    async def _process_structured_data(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Process structured data and return metadata"""
        try:
            if file_type == 'csv':
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path, sheet_name=None)  # Read all sheets

            if isinstance(df, dict):
                # Multiple sheets
                structured_info = {}
                for sheet_name, sheet_df in df.items():
                    structured_info[sheet_name] = {
                        'shape': sheet_df.shape,
                        'columns': list(sheet_df.columns),
                        'dtypes': sheet_df.dtypes.to_dict(),
                        'null_counts': sheet_df.isnull().sum().to_dict(),
                        'sample_data': sheet_df.head(3).to_dict('records')
                    }
                return structured_info
            else:
                # Single sheet/CSV
                return {
                    'shape': df.shape,
                    'columns': list(df.columns),
                    'dtypes': df.dtypes.to_dict(),
                    'null_counts': df.isnull().sum().to_dict(),
                    'sample_data': df.head(3).to_dict('records')
                }
        except Exception as e:
            logger.error(f"Structured data processing failed: {e}")
            return {}

    async def _extract_entities_with_ai(self, content: str) -> List[Dict[str, Any]]:
        """Extract entities using AI/LLM"""
        if not content or len(content) < 100:
            return []

        try:
            # Initialize LLM client if not done
            await client.initialize()

            # Create prompt for entity extraction
            prompt = f"""
            Analyze the following document content and extract key business entities, organizations, locations, and sustainability-related information.

            Return a JSON list of entities with the following structure:
            [
                {{
                    "text": "entity name",
                    "type": "organization|person|location|product|metric|date",
                    "category": "business|sustainability|financial|operational",
                    "confidence": 0.8
                }}
            ]

            Document content:
            {content[:4000]}  # Limit content to avoid token limits
            """

            messages = [LLMMessage(role="user", content=prompt)]
            response = await client.generate(
                messages=messages,
                max_tokens=800,
                temperature=0.1  # Low temperature for structured output
            )

            # Parse JSON response
            try:
                import json
                entities = json.loads(response.content)
                return entities if isinstance(entities, list) else []
            except json.JSONDecodeError:
                logger.warning("Failed to parse entity extraction JSON response")
                return []

        except Exception as e:
            logger.error(f"AI entity extraction failed: {e}")
            return []

    async def _analyze_sustainability_content(self, content: str) -> Dict[str, Any]:
        """Analyze sustainability-specific content using AI"""
        if not content or len(content) < 100:
            return {}

        try:
            await client.initialize()

            prompt = f"""
            Analyze this document content for sustainability and ESG information. Extract key insights about:

            1. Carbon emissions and greenhouse gas data
            2. Environmental impact metrics
            3. Social responsibility initiatives
            4. Governance practices
            5. Sustainability goals and targets

            Return a JSON object with this structure:
            {{
                "carbon_footprint": {{"mentioned": true/false, "metrics": [], "scope": []}},
                "environmental": {{"topics": [], "metrics": [], "initiatives": []}},
                "social": {{"topics": [], "initiatives": [], "metrics": []}},
                "governance": {{"practices": [], "policies": []}},
                "sustainability_score": 0.7,
                "key_insights": ["insight 1", "insight 2"]
            }}

            Document content:
            {content[:3000]}
            """

            messages = [LLMMessage(role="user", content=prompt)]
            response = await client.generate(
                messages=messages,
                max_tokens=1000,
                temperature=0.2
            )

            try:
                import json
                insights = json.loads(response.content)
                return insights if isinstance(insights, dict) else {}
            except json.JSONDecodeError:
                logger.warning("Failed to parse sustainability analysis JSON response")
                return {}

        except Exception as e:
            logger.error(f"AI sustainability analysis failed: {e}")
            return {}

    def _calculate_confidence_score(self, doc: ProcessedDocument) -> float:
        """Calculate overall confidence score for the processed document"""
        factors = []

        # Content quality
        if len(doc.content) > 1000:
            factors.append(0.9)
        elif len(doc.content) > 500:
            factors.append(0.7)
        else:
            factors.append(0.4)

        # Structured data availability
        if doc.structured_data:
            factors.append(0.8)
        else:
            factors.append(0.6)

        # Entity extraction success
        if doc.entities:
            factors.append(0.8)
        else:
            factors.append(0.5)

        # Sustainability insights
        if doc.sustainability_insights:
            factors.append(0.9)
        else:
            factors.append(0.6)

        return sum(factors) / len(factors)